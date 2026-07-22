# Script to generate project documentation in Word format.
import os
import glob
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

output_dir = r"/mnt/c/Users/RYZEN/Desktop/projects/fast_api_generator_backend/outputs/4cb7114e-dc00-4a30-bd4f-f3fb96d29919"

skip_dirs = [".venv", "venv", "__pycache__", ".git", "node_modules", ".code_index"]
generated_files = []
for f in sorted(Path(output_dir).rglob("*")):
    if f.is_file() and not any(p in f.parts for p in skip_dirs):
        generated_files.append(str(f.relative_to(output_dir)))

readme_text = ""
readme_path = os.path.join(output_dir, "README.md")
if os.path.isfile(readme_path):
    with open(readme_path, errors="replace") as f:
        readme_text = f.read()

logs_text = ""
for lf in sorted(glob.glob(os.path.join(output_dir, "server_logs*.md")), key=lambda p: (len(p), p)):
    with open(lf, errors="replace") as f:
        logs_text += "\n\n=== " + os.path.basename(lf) + " ===\n" + f.read()


def add_h1(doc, text):
    p = doc.add_heading(text, 1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


project_name = Path(output_dir).name

doc = Document()
title = doc.add_heading("Project Documentation", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

add_h1(doc, "Generated Files")
doc.add_paragraph("All files generated for this project:")
for fp in generated_files:
    add_bullet(doc, fp)

if readme_text.strip():
    add_h1(doc, "Architecture and Setup")
    para = doc.add_paragraph()
    run = para.add_run(readme_text.strip())
    run.font.size = Pt(9)

if logs_text.strip():
    add_h1(doc, "API Test Results")
    para = doc.add_paragraph()
    run = para.add_run(logs_text.strip())
    run.font.name = "Courier New"
    run.font.size = Pt(8)

doc.save("project_report.docx")
print("Saved: project_report.docx")
